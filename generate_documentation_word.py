"""
Script para generar documento Word profesional desde documentación markdown
Incluye logo en header y estilos profesionales para presentación a clientes
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("=" * 60)
    print("Error: python-docx no está instalado")
    print("=" * 60)
    print("Para instalarlo ejecuta:")
    print("  pip install python-docx")
    print()
    print("O con --user:")
    print("  pip install --user python-docx")
    print("=" * 60)
    sys.exit(1)

try:
    import markdown
except ImportError:
    print("Instalando markdown...")
    os.system(f"{sys.executable} -m pip install markdown")
    import markdown


def add_header_with_logo(doc, logo_path):
    """Agrega header con logo al documento"""
    try:
        # Obtener la sección
        section = doc.sections[0]
        
        # Crear header
        header = section.header
        
        # Limpiar header existente
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Crear párrafo en el header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar logo si existe
        if os.path.exists(logo_path):
            try:
                run = header_para.add_run()
                run.add_picture(logo_path, width=Inches(3))
                # Agregar espacio después del logo
                header_para.add_run().add_break()
                header_para.add_run().add_break()
            except Exception as e:
                print(f"Advertencia: No se pudo agregar el logo al header: {e}")
        
    except Exception as e:
        print(f"Advertencia: No se pudo crear el header: {e}")


def add_footer_with_page_number(doc):
    """Agrega footer con número de página"""
    try:
        section = doc.sections[0]
        footer = section.footer
        
        # Limpiar footer existente
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # Crear párrafo en el footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar número de página
        run = footer_para.add_run()
        run.text = "Página "
        
        # Campo de página actual
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        run = footer_para.add_run()
        run.text = " de "
        
        # Campo de total de páginas
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'NUMPAGES'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # Estilo del footer
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(148, 163, 184)  # #94a3b8
        
    except Exception as e:
        print(f"Advertencia: No se pudo crear el footer: {e}")


def convert_markdown_to_word(markdown_text, doc):
    """Convierte markdown a elementos de Word"""
    # Convertir markdown a HTML
    md = markdown.Markdown(extensions=['extra', 'tables', 'fenced_code', 'nl2br'])
    html = md.convert(markdown_text)
    
    # Usar BeautifulSoup si está disponible
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'li', 'ul', 'ol', 'hr', 'code', 'pre']):
            tag_name = element.name
            
            if tag_name == 'h1':
                p = doc.add_paragraph()
                p.style = 'Heading 1'
                run = p.add_run(element.get_text())
                run.font.color.rgb = RGBColor(30, 41, 59)  # #1e293b
                run.font.size = Pt(24)
                run.bold = True
                doc.add_paragraph()  # Espacio después
                
            elif tag_name == 'h2':
                p = doc.add_paragraph()
                p.style = 'Heading 2'
                run = p.add_run(element.get_text())
                run.font.color.rgb = RGBColor(51, 65, 85)  # #334155
                run.font.size = Pt(18)
                run.bold = True
                doc.add_paragraph()  # Espacio después
                
            elif tag_name == 'h3':
                p = doc.add_paragraph()
                p.style = 'Heading 3'
                run = p.add_run(element.get_text())
                run.font.color.rgb = RGBColor(71, 85, 105)  # #475569
                run.font.size = Pt(14)
                run.bold = True
                doc.add_paragraph()  # Espacio después
                
            elif tag_name == 'h4':
                p = doc.add_paragraph()
                p.style = 'Heading 4'
                run = p.add_run(element.get_text())
                run.font.color.rgb = RGBColor(100, 116, 139)  # #64748b
                run.font.size = Pt(12)
                run.bold = True
                doc.add_paragraph()  # Espacio después
                
            elif tag_name == 'p':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Procesar contenido del párrafo manteniendo formato
                for content in element.contents:
                    if hasattr(content, 'name'):
                        if content.name == 'strong' or content.name == 'b':
                            run = p.add_run(content.get_text())
                            run.bold = True
                        elif content.name == 'em' or content.name == 'i':
                            run = p.add_run(content.get_text())
                            run.italic = True
                        elif content.name == 'code':
                            run = p.add_run(content.get_text())
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(225, 29, 72)  # #e11d48
                        else:
                            p.add_run(content.get_text() if hasattr(content, 'get_text') else str(content))
                    else:
                        text = str(content).strip()
                        if text:
                            p.add_run(text)
                
                # Si el párrafo está vacío, usar texto completo
                if not p.runs:
                    p.add_run(element.get_text())
                
            elif tag_name == 'li':
                p = doc.add_paragraph(style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Procesar contenido de la lista
                for content in element.contents:
                    if hasattr(content, 'name'):
                        if content.name == 'strong' or content.name == 'b':
                            run = p.add_run(content.get_text())
                            run.bold = True
                        elif content.name == 'em' or content.name == 'i':
                            run = p.add_run(content.get_text())
                            run.italic = True
                        elif content.name == 'code':
                            run = p.add_run(content.get_text())
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(225, 29, 72)
                        else:
                            p.add_run(content.get_text() if hasattr(content, 'get_text') else str(content))
                    else:
                        text = str(content).strip()
                        if text:
                            p.add_run(text)
                
                if not p.runs:
                    p.add_run(element.get_text())
                    
            elif tag_name == 'hr':
                p = doc.add_paragraph()
                p_format = p.paragraph_format
                p_format.space_before = Pt(12)
                p_format.space_after = Pt(12)
                
            elif tag_name == 'pre' or tag_name == 'code':
                # Bloque de código
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85)  # #334155
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)
                doc.add_paragraph()  # Espacio después
    
    except ImportError:
        # Fallback sin BeautifulSoup - procesamiento simple
        import re
        
        # Normalizar HTML
        html = re.sub(r'<strong>', '<b>', html, flags=re.IGNORECASE)
        html = re.sub(r'</strong>', '</b>', html, flags=re.IGNORECASE)
        html = re.sub(r'<em>', '<i>', html, flags=re.IGNORECASE)
        html = re.sub(r'</em>', '</i>', html, flags=re.IGNORECASE)
        
        # Extraer títulos
        for match in re.finditer(r'<h([1-4])[^>]*>(.*?)</h[1-4]>', html, re.DOTALL | re.IGNORECASE):
            level = int(match.group(1))
            text = re.sub(r'<[^>]+>', '', match.group(2)).strip()
            if text:
                p = doc.add_paragraph()
                if level == 1:
                    p.style = 'Heading 1'
                    run = p.add_run(text)
                    run.font.size = Pt(24)
                elif level == 2:
                    p.style = 'Heading 2'
                    run = p.add_run(text)
                    run.font.size = Pt(18)
                elif level == 3:
                    p.style = 'Heading 3'
                    run = p.add_run(text)
                    run.font.size = Pt(14)
                elif level == 4:
                    p.style = 'Heading 4'
                    run = p.add_run(text)
                    run.font.size = Pt(12)
                run.bold = True
                doc.add_paragraph()
        
        # Extraer párrafos
        for match in re.finditer(r'<p[^>]*>(.*?)</p>', html, re.DOTALL | re.IGNORECASE):
            para_text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()
        
        # Extraer listas
        for match in re.finditer(r'<li[^>]*>(.*?)</li>', html, re.DOTALL | re.IGNORECASE):
            li_text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
            if li_text:
                doc.add_paragraph(li_text, style='List Bullet')


def generate_word_from_markdown(md_file_path, output_docx_path, logo_path):
    """
    Convierte un archivo markdown a Word con logo en header
    
    Args:
        md_file_path: Ruta al archivo markdown
        output_docx_path: Ruta donde se guardará el documento Word
        logo_path: Ruta al logo para el header
    """
    
    # Leer el archivo markdown
    print(f"Leyendo archivo markdown: {md_file_path}")
    with open(md_file_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Crear documento Word
    print(f"Generando documento Word: {output_docx_path}")
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Agregar header con logo
    print("Agregando header con logo...")
    add_header_with_logo(doc, logo_path)
    
    # Agregar footer con número de página
    print("Agregando footer...")
    add_footer_with_page_number(doc)
    
    # Convertir markdown a Word
    print("Convirtiendo markdown a Word...")
    convert_markdown_to_word(markdown_content, doc)
    
    # Guardar documento
    print("Guardando documento...")
    output_dir = os.path.dirname(output_docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc.save(output_docx_path)
    
    print(f"[OK] Documento Word generado exitosamente en: {output_docx_path}")
    return True


def main():
    """Función principal"""
    # Rutas de archivos
    base_dir = Path(__file__).parent
    md_file = base_dir / "documents" / "PMS_Funcionalidades_y_Modulos.md"
    logo_file = base_dir / "frontend" / "src" / "assets" / "img" / "logo_complet_black_transparent.png"
    output_docx = base_dir / "documents" / "AlojaSys_Documentacion.docx"
    
    # Verificar que los archivos existen
    if not md_file.exists():
        print(f"Error: No se encontró el archivo markdown: {md_file}")
        return
    
    if not logo_file.exists():
        print(f"Error: No se encontró el logo: {logo_file}")
        return
    
    print("=" * 60)
    print("Generador de Word - Documentación AlojaSys")
    print("=" * 60)
    print()
    
    # Generar Word
    try:
        generate_word_from_markdown(
            str(md_file),
            str(output_docx),
            str(logo_file)
        )
        print()
        print("=" * 60)
        print("¡Documento Word generado exitosamente!")
        print(f"Ubicación: {output_docx}")
        print()
        print("Puedes abrirlo en Microsoft Word y guardarlo como PDF")
        print("=" * 60)
    except Exception as e:
        print()
        print("=" * 60)
        print(f"Error al generar el documento Word: {str(e)}")
        import traceback
        traceback.print_exc()
        print("=" * 60)
        sys.exit(1)


if __name__ == "__main__":
    main()

