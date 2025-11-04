"""
Script para generar documento Word resumido para clientes
Versión ejecutiva y comercial de la documentación AlojaSys
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx no está instalado")
    print("Instálalo con: pip install python-docx")
    sys.exit(1)


def add_header_with_logo(doc, logo_path):
    """Agrega header con logo al documento"""
    try:
        section = doc.sections[0]
        header = section.header
        
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(logo_path):
            try:
                run = header_para.add_run()
                run.add_picture(logo_path, width=Inches(3))
                header_para.add_run().add_break()
                header_para.add_run().add_break()
            except Exception as e:
                print(f"Advertencia: No se pudo agregar el logo: {e}")
        
    except Exception as e:
        print(f"Advertencia: No se pudo crear el header: {e}")


def add_footer_with_page_number(doc):
    """Agrega footer con número de página"""
    try:
        section = doc.sections[0]
        footer = section.footer
        
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = footer_para.add_run()
        run.text = "Página "
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        run = footer_para.add_run()
        run.text = " de "
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'NUMPAGES'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(148, 163, 184)
        
    except Exception as e:
        print(f"Advertencia: No se pudo crear el footer: {e}")


def generate_resumen_cliente(doc, logo_path):
    """Genera el contenido del resumen para clientes"""
    
    # PORTADA / TÍTULO
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AlojaSys")
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(30, 41, 59)
    title_run.bold = True
    
    doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Sistema de Gestión Hotelera")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    date_run = date_para.add_run(f"Documento de Presentación - {datetime.now().strftime('%B %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()
    
    # INTRODUCCIÓN
    heading1 = doc.add_heading("¿Por qué AlojaSys?", 1)
    heading1.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "AlojaSys es una solución integral para la gestión hotelera que automatiza y simplifica "
        "todos los aspectos administrativos de su hotel. Desde la reserva hasta el check-out, "
        "desde los pagos hasta la facturación, todo en un solo sistema fácil de usar."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # MÓDULOS PRINCIPALES
    heading2 = doc.add_heading("Funcionalidades Principales", 1)
    heading2.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    # 1. Gestión de Reservas
    doc.add_heading("1. Gestión de Reservas", 2)
    doc.add_paragraph(
        "Administre todas sus reservas desde un calendario visual intuitivo. El sistema maneja automáticamente "
        "los estados (pendiente, confirmada, check-in, check-out) y puede realizar check-in/out automáticos "
        "según la configuración del hotel.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Visualización clara de disponibilidad, prevención de doble reserva y gestión completa del ciclo de vida de cada reserva.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 2. Gestión de Habitaciones
    doc.add_heading("2. Gestión de Habitaciones", 2)
    doc.add_paragraph(
        "Configure diferentes tipos de habitaciones con precios flexibles, capacidad y estados en tiempo real.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Sistema de precios base más extras por huéspedes adicionales, control de disponibilidad y gestión de mantenimiento.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 3. Sistema de Pagos
    doc.add_heading("3. Sistema de Pagos Integral", 2)
    doc.add_paragraph(
        "Integración con múltiples métodos de pago: Mercado Pago, transferencias bancarias, efectivo y más.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Sistema de señas (pagos parciales) con generación automática de comprobantes.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reconocimiento automático de transferencias bancarias mediante OCR para identificación instantánea de pagos.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Conciliación bancaria automática con extractos para control total de ingresos.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Procesamiento automático de reembolsos según políticas configuradas.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 4. Facturación
    doc.add_heading("4. Facturación Electrónica", 2)
    doc.add_paragraph(
        "Facturación electrónica AFIP completa para cumplimiento normativo argentino.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Generación automática de comprobantes de señas y pagos parciales.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Integración con sistemas contables para exportación de datos.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 5. Políticas y Tarifas
    doc.add_heading("5. Políticas y Tarifas Flexibles", 2)
    doc.add_paragraph(
        "Configuración de políticas de cancelación y devolución personalizables.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Sistema de tarifas dinámicas con diferentes precios según temporada, días de semana o eventos.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Gestión de penalidades por no-show automáticas.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 6. Dashboard y Reportes
    doc.add_heading("6. Dashboard y Reportes", 2)
    doc.add_paragraph(
        "Panel de control con métricas clave: ocupación, ingresos, reservas pendientes y más.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reportes detallados de ingresos, ocupación, huéspedes y análisis de rentabilidad.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Visualización de datos en tiempo real para toma de decisiones informadas.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 7. Calendario Visual
    doc.add_heading("7. Calendario Interactivo", 2)
    doc.add_paragraph(
        "Vista mensual y semanal de todas las reservas con información detallada al pasar el mouse.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Filtros por habitación, estado y fechas para fácil navegación.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Gestión rápida de reservas directamente desde el calendario.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 8. Gestión de Usuarios
    doc.add_heading("8. Control de Accesos", 2)
    doc.add_paragraph(
        "Sistema de roles y permisos para personalizar el acceso según el cargo del empleado.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Roles predefinidos: Administrador, Recepción, Contador, con permisos específicos.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Gestión de múltiples hoteles desde una sola cuenta con control granular.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 9. Notificaciones
    doc.add_heading("9. Sistema de Notificaciones", 2)
    doc.add_paragraph(
        "Alertas automáticas sobre eventos importantes: nuevas reservas, pagos recibidos, check-ins pendientes.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Notificaciones por email y en el sistema para mantenerlo siempre informado.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # 10. Integraciones
    doc.add_heading("10. Integraciones con OTAs", 2)
    doc.add_paragraph(
        "Conexión con principales plataformas de reservas online (Booking, Airbnb, Expedia, etc.).",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Sincronización automática de disponibilidad y precios para maximizar ocupación.",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # BENEFICIOS
    doc.add_page_break()
    heading3 = doc.add_heading("Beneficios Clave", 1)
    heading3.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    beneficios = [
        ("Ahorro de Tiempo", "Automatización de procesos manuales reduce el trabajo administrativo en un 70%"),
        ("Mayor Precisión", "Eliminación de errores humanos en cálculos, reservas y facturación"),
        ("Control Total", "Visibilidad completa de su negocio en tiempo real con reportes detallados"),
        ("Escalabilidad", "Administre un hotel o múltiples propiedades desde la misma plataforma"),
        ("Cumplimiento Legal", "Facturación electrónica AFIP integrada para cumplimiento normativo"),
        ("Mejor Experiencia", "Interfaz intuitiva que requiere mínimo entrenamiento para su personal"),
        ("Automatización 24/7", "Procesos como check-out, reembolsos y conciliación funcionan automáticamente"),
        ("Maximización de Ingresos", "Herramientas para optimizar precios, ocupación y gestión de reservas"),
    ]
    
    for titulo, descripcion in beneficios:
        doc.add_heading(titulo, 2)
        para = doc.add_paragraph(descripcion)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.runs[0].font.size = Pt(11)
        doc.add_paragraph()
    
    # CASOS DE USO
    doc.add_page_break()
    heading4 = doc.add_heading("Casos de Uso Comunes", 1)
    heading4.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    casos = [
        (
            "Recepción de Reserva",
            "Un huésped llama para reservar. El recepcionista consulta disponibilidad en el calendario, "
            "crea la reserva, procesa la seña por Mercado Pago y el sistema genera automáticamente el comprobante. "
            "La habitación se marca como reservada y se envía confirmación por email."
        ),
        (
            "Check-in Automático",
            "El día de llegada, el sistema marca automáticamente la reserva como check-in a las 15:00 (según configuración). "
            "La habitación cambia a 'ocupada' y se notifica al personal. No requiere intervención manual."
        ),
        (
            "Procesamiento de Pago",
            "El huésped realiza una transferencia bancaria. El sistema reconoce automáticamente el pago mediante OCR, "
            "lo asocia a la reserva correspondiente y actualiza el estado de pago. La conciliación bancaria se realiza automáticamente."
        ),
        (
            "Reembolso Automático",
            "Un huésped cancela según la política configurada. El sistema calcula automáticamente el monto a reembolsar, "
            "procesa el reembolso según la política y notifica al huésped. Todo sin intervención manual."
        ),
        (
            "Análisis de Negocio",
            "El gerente accede al dashboard y visualiza métricas clave: ocupación del 85%, ingresos del mes, "
            "reservas pendientes y comparativa con meses anteriores. Toma decisiones informadas sobre precios y promociones."
        ),
    ]
    
    for titulo, descripcion in casos:
        doc.add_heading(titulo, 2)
        para = doc.add_paragraph(descripcion)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.runs[0].font.size = Pt(11)
        doc.add_paragraph()
    
    # CARACTERÍSTICAS TÉCNICAS (Resumidas)
    doc.add_page_break()
    heading5 = doc.add_heading("Características Técnicas", 1)
    heading5.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    doc.add_paragraph(
        "• Plataforma web accesible desde cualquier dispositivo con conexión a internet",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Almacenamiento seguro en la nube con respaldos automáticos",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Interfaz responsive para uso en computadoras, tablets y móviles",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Integración con sistemas de pago y facturación electrónica",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Soporte multi-hotel para cadenas o grupos hoteleros",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Actualizaciones continuas sin interrupciones del servicio",
        style='List Bullet'
    )
    doc.add_paragraph()
    
    # CIERRE
    doc.add_page_break()
    heading6 = doc.add_heading("Conclusión", 1)
    heading6.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    conclusion = doc.add_paragraph()
    conclusion_run = conclusion.add_run(
        "AlojaSys es la solución moderna que su hotel necesita para optimizar operaciones, "
        "reducir costos administrativos y mejorar la experiencia tanto para su personal como para sus huéspedes. "
        "Con automatización inteligente, control total y cumplimiento normativo integrado, "
        "podrá enfocarse en lo que realmente importa: brindar excelente servicio a sus huéspedes."
    )
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    conclusion_run.font.size = Pt(12)
    conclusion_run.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    contacto_para = doc.add_paragraph()
    contacto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contacto_run = contacto_para.add_run("Para más información, contacte a nuestro equipo comercial")
    contacto_run.font.size = Pt(11)
    contacto_run.italic = True
    contacto_run.font.color.rgb = RGBColor(100, 116, 139)


def generate_word_resumen(output_docx_path, logo_path):
    """Genera el documento Word resumido"""
    
    print(f"Generando documento Word resumido: {output_docx_path}")
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Agregar header con logo
    print("Agregando header con logo...")
    add_header_with_logo(doc, logo_path)
    
    # Agregar footer con número de página
    print("Agregando footer...")
    add_footer_with_page_number(doc)
    
    # Generar contenido
    print("Generando contenido del resumen...")
    generate_resumen_cliente(doc, logo_path)
    
    # Guardar documento
    print("Guardando documento...")
    output_dir = os.path.dirname(output_docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc.save(output_docx_path)
    
    print(f"[OK] Documento Word generado exitosamente en: {output_docx_path}")
    return True


def main():
    """Función principal"""
    base_dir = Path(__file__).parent
    logo_file = base_dir / "frontend" / "src" / "assets" / "img" / "logo_complet_black_transparent.png"
    output_docx = base_dir / "documents" / "AlojaSys_Resumen_Cliente.docx"
    
    if not logo_file.exists():
        print(f"Error: No se encontró el logo: {logo_file}")
        return
    
    print("=" * 60)
    print("Generador de Resumen - AlojaSys para Clientes")
    print("=" * 60)
    print()
    
    try:
        generate_word_resumen(str(output_docx), str(logo_file))
        print()
        print("=" * 60)
        print("[OK] Documento Word generado exitosamente!")
        print(f"Ubicacion: {output_docx}")
        print()
        print("Puedes abrirlo en Microsoft Word y guardarlo como PDF")
        print("=" * 60)
    except Exception as e:
        print()
        print("=" * 60)
        print(f"Error al generar el documento: {str(e)}")
        import traceback
        traceback.print_exc()
        print("=" * 60)
        sys.exit(1)


if __name__ == "__main__":
    main()

